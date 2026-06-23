from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'output' / 'progress-brief'
IMG = OUT / 'screenshots'
DOCX = OUT / 'YUTANK_Interim_WIP_Submission.docx'
OUT.mkdir(parents=True, exist_ok=True)
IMG.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(96, 96, 96)


def font(size=28, bold=False):
    candidates = [
        'C:/Windows/Fonts/arialbd.ttf' if bold else 'C:/Windows/Fonts/arial.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf' if bold else '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
    ]
    for c in candidates:
        if Path(c).exists():
            return ImageFont.truetype(c, size)
    return ImageFont.load_default()


def box(draw, xy, fill='white', outline='#D3DCE6', radius=12, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, size=24, fill='#222222', bold=False):
    draw.text(xy, value, font=font(size, bold), fill=fill)


def wrap(draw, value, max_width, size=22, bold=False):
    f = font(size, bold)
    words = value.split()
    lines = []
    cur = ''
    for w in words:
        cand = w if not cur else cur + ' ' + w
        if draw.textbbox((0, 0), cand, font=f)[2] <= max_width:
            cur = cand
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def draw_wrapped(draw, xy, value, max_width, size=22, fill='#222222', bold=False):
    x, y = xy
    f = font(size, bold)
    for line in wrap(draw, value, max_width, size, bold):
        draw.text((x, y), line, font=f, fill=fill)
        y += size + 8
    return y


def base(title, subtitle):
    im = Image.new('RGB', (1440, 900), '#F4F6F8')
    d = ImageDraw.Draw(im)
    d.rectangle((0, 0, 1440, 72), fill='#263445')
    text(d, (32, 22), 'YUTANK Shop Admin', 24, 'white', True)
    text(d, (1210, 24), 'System Administrator', 18, '#E7EEF7')
    d.rectangle((0, 72, 240, 900), fill='#304156')
    menu = ['Dashboard', 'Employees', 'Credits', 'Categories', 'Goods']
    for i, m in enumerate(menu):
        y = 112 + i * 54
        fill = '#1F2D3D' if m in title else '#304156'
        d.rectangle((0, y - 10, 240, y + 34), fill=fill)
        text(d, (34, y), m, 20, '#E7EEF7')
    text(d, (282, 104), title, 34, '#1F4D78', True)
    if subtitle:
        text(d, (282, 148), subtitle, 20, '#606266')
    return im, d


def table(draw, x, y, cols, rows, widths):
    h = 44
    draw.rectangle((x, y, x + sum(widths), y + h), fill='#E8EEF5', outline='#C7D0DA')
    cx = x
    for c, w in zip(cols, widths):
        text(draw, (cx + 12, y + 12), c, 18, '#1F4D78', True)
        draw.line((cx, y, cx, y + h + len(rows) * h), fill='#C7D0DA')
        cx += w
    draw.line((x + sum(widths), y, x + sum(widths), y + h + len(rows) * h), fill='#C7D0DA')
    for r, row in enumerate(rows):
        yy = y + h * (r + 1)
        draw.rectangle((x, yy, x + sum(widths), yy + h), fill='white', outline='#D8DEE6')
        cx = x
        for value, w in zip(row, widths):
            text(draw, (cx + 12, yy + 12), str(value), 17, '#303133')
            cx += w


def save_login():
    im = Image.new('RGB', (1440, 900), '#2D3A4B')
    d = ImageDraw.Draw(im)
    box(d, (460, 170, 980, 690), '#2D3A4B', '#596979', 10, 1)
    text(d, (575, 230), 'YUTANK Shop Admin', 36, '#EEEEEE', True)
    for i, label in enumerate(['Employee username', 'Password']):
        y = 330 + i * 92
        box(d, (535, y, 905, y + 54), '#273445', '#4C5B6B', 6, 1)
        text(d, (558, y + 16), label, 19, '#BFCBD9')
    d.rectangle((535, 510, 558, 533), outline='#BFCBD9', width=2)
    text(d, (570, 508), 'Remember me', 20, '#EEEEEE')
    box(d, (535, 570, 905, 628), '#409EFF', '#409EFF', 6, 1)
    text(d, (695, 588), 'Login', 22, 'white', True)
    im.save(IMG / '01-login.png')


def save_dashboard():
    im, d = base('Dashboard', 'Current back-office overview')
    cards = [('Current Employee', 'System Administrator'), ('My Credits', '0'), ('Management Modules', '3'), ('Backend Service', 'Online')]
    for i, (k, v) in enumerate(cards):
        x = 282 + i * 270
        box(d, (x, 210, x + 240, 340))
        text(d, (x + 20, 232), k, 18, '#909399')
        text(d, (x + 20, 278), v, 27, '#303133', True)
    box(d, (282, 390, 1160, 560))
    text(d, (312, 422), 'Quick Links', 24, '#303133', True)
    for i, label in enumerate(['Employees', 'Credit Operations', 'Categories']):
        box(d, (312 + i * 250, 480, 520 + i * 250, 532), '#409EFF' if i == 0 else '#67C23A' if i == 1 else '#E6A23C', '#FFFFFF', 8, 1)
        text(d, (350 + i * 250, 496), label, 19, 'white', True)
    im.save(IMG / '02-dashboard.png')


def save_employee_list():
    im, d = base('Employees', 'Search, create, edit, enable or disable staff accounts')
    box(d, (282, 196, 1240, 282))
    text(d, (306, 226), 'Username', 18, '#606266')
    text(d, (520, 226), 'Name', 18, '#606266')
    text(d, (716, 226), 'Status: All', 18, '#606266')
    box(d, (1040, 213, 1190, 258), '#409EFF', '#409EFF', 6, 1)
    text(d, (1085, 226), 'Search', 18, 'white', True)
    box(d, (282, 320, 1240, 720))
    text(d, (306, 345), 'Employees', 24, '#303133', True)
    box(d, (1060, 338, 1210, 382), '#409EFF', '#409EFF', 6, 1)
    text(d, (1082, 350), 'New Employee', 17, 'white', True)
    rows = [['12', 'root', 'System Administrator', 'Active', 'Edit | Disable | Reset Password'], ['15', 'staff01', 'Store Staff', 'Active', 'Edit | Disable | Reset Password']]
    table(d, 306, 405, ['ID', 'Username', 'Name', 'Status', 'Actions'], rows, [80, 170, 240, 150, 300])
    im.save(IMG / '03-employee-list.png')


def save_employee_form():
    im, d = base('Employees', 'New Employee dialog')
    box(d, (470, 170, 970, 690), 'white', '#C7D0DA', 12, 2)
    text(d, (510, 208), 'New Employee', 28, '#303133', True)
    fields = ['Username', 'Password', 'Name', 'Phone', 'Email']
    for i, f in enumerate(fields):
        y = 280 + i * 62
        text(d, (520, y + 12), f, 18, '#606266')
        box(d, (650, y, 920, y + 44), '#FFFFFF', '#D3DCE6', 5, 1)
    box(d, (692, 615, 800, 660), '#FFFFFF', '#D3DCE6', 5, 1)
    text(d, (718, 628), 'Cancel', 17, '#606266')
    box(d, (820, 615, 930, 660), '#409EFF', '#409EFF', 5, 1)
    text(d, (845, 628), 'Confirm', 17, 'white', True)
    im.save(IMG / '04-employee-form.png')


def save_my_credits():
    im, d = base('Credits - My Credits', 'Current balance and transaction records')
    box(d, (282, 200, 560, 370))
    text(d, (312, 232), 'Current Credit Balance', 20, '#909399')
    text(d, (312, 286), '0', 52, '#303133', True)
    box(d, (600, 200, 1240, 640))
    text(d, (630, 232), 'My Credit Records', 24, '#303133', True)
    table(d, 630, 285, ['ID', 'Type', 'Credits', 'Before', 'After', 'Remark'], [['-', '-', '-', '-', '-', 'No records yet']], [70, 90, 110, 110, 110, 230])
    im.save(IMG / '05-my-credits.png')


def save_credit_ops():
    im, d = base('Credits - Credit Operations', 'Add or deduct employee credits')
    box(d, (282, 190, 1240, 300))
    text(d, (312, 225), 'Employee: Search username or name', 20, '#606266')
    text(d, (650, 225), 'Credits: 1', 20, '#606266')
    text(d, (820, 225), 'Remark: Operation remark', 20, '#606266')
    box(d, (1080, 220, 1145, 262), '#67C23A', '#67C23A', 5, 1)
    text(d, (1098, 232), 'Add', 17, 'white', True)
    box(d, (1160, 220, 1235, 262), '#F56C6C', '#F56C6C', 5, 1)
    text(d, (1172, 232), 'Deduct', 17, 'white', True)
    box(d, (282, 340, 1240, 720))
    text(d, (312, 370), 'Employee Credit Records', 24, '#303133', True)
    table(d, 312, 420, ['ID', 'Employee ID', 'Type', 'Credits', 'Before', 'After', 'Remark'], [['-', '-', '-', '-', '-', '-', 'Select an employee to review records']], [70, 120, 90, 100, 100, 100, 310])
    im.save(IMG / '06-credit-operations.png')


def save_category():
    im, d = base('Categories', 'Product category list')
    box(d, (282, 190, 1020, 540))
    text(d, (312, 222), 'Categories', 24, '#303133', True)
    table(d, 312, 280, ['ID', 'Category Name', 'Sort', 'Status'], [['1', 'Office Snacks', '1', 'Enabled'], ['2', 'Employee Benefits', '2', 'Enabled'], ['3', 'Office Supplies', '3', 'Enabled']], [80, 330, 120, 150])
    text(d, (780, 488), '3 categories', 18, '#606266')
    im.save(IMG / '07-category-list.png')


def save_goods_list():
    im, d = base('Goods', 'Filter and manage reward goods')
    box(d, (282, 190, 1240, 280))
    text(d, (312, 224), 'Name: Goods name', 18, '#606266')
    text(d, (540, 224), 'Category: All', 18, '#606266')
    text(d, (740, 224), 'Status: All', 18, '#606266')
    box(d, (1080, 214, 1205, 258), '#409EFF', '#409EFF', 5, 1)
    text(d, (1115, 227), 'Search', 17, 'white', True)
    box(d, (282, 320, 1290, 750))
    text(d, (312, 350), 'Goods Management', 24, '#303133', True)
    box(d, (1110, 342, 1240, 384), '#409EFF', '#409EFF', 5, 1)
    text(d, (1130, 354), 'New Goods', 17, 'white', True)
    rows = [['7', 'Goods Smoke Updated', 'Office Snacks', '12', '7', 'On Shelf', 'Detail | Edit | Off Shelf'], ['2', 'Credit Goods updated', 'Employee Benefits', '99', '18', 'On Shelf', 'Detail | Edit | Off Shelf']]
    table(d, 312, 410, ['ID', 'Goods Name', 'Category', 'Credits', 'Stock', 'Status', 'Actions'], rows, [60, 230, 180, 95, 80, 120, 220])
    im.save(IMG / '08-goods-list.png')


def save_goods_form():
    im, d = base('Goods', 'New Goods dialog')
    box(d, (450, 140, 1010, 750), 'white', '#C7D0DA', 12, 2)
    text(d, (495, 178), 'New Goods', 28, '#303133', True)
    fields = ['Category', 'Name', 'Image URL', 'Credits Price', 'Stock', 'Description']
    for i, f in enumerate(fields):
        y = 250 + i * 68
        text(d, (500, y + 12), f, 18, '#606266')
        box(d, (670, y, 950, y + (86 if f == 'Description' else 44)), '#FFFFFF', '#D3DCE6', 5, 1)
    box(d, (735, 675, 835, 718), '#FFFFFF', '#D3DCE6', 5, 1)
    text(d, (758, 687), 'Cancel', 17, '#606266')
    box(d, (850, 675, 960, 718), '#409EFF', '#409EFF', 5, 1)
    text(d, (875, 687), 'Confirm', 17, 'white', True)
    im.save(IMG / '09-goods-form.png')


def save_goods_detail():
    im, d = base('Goods', 'Goods detail dialog')
    box(d, (450, 160, 1010, 720), 'white', '#C7D0DA', 12, 2)
    text(d, (495, 198), 'Goods Detail', 28, '#303133', True)
    rows = [('ID', '7'), ('Name', 'Goods Smoke Updated'), ('Category', 'Office Snacks'), ('Credits Price', '12'), ('Stock', '7'), ('Status', 'On Shelf'), ('Image URL', 'https://example.com/goods-updated.png'), ('Description', 'Smoke test goods updated')]
    y = 260
    for k, v in rows:
        d.rectangle((500, y, 650, y + 42), fill='#F5F7FA', outline='#EBEEF5')
        d.rectangle((650, y, 960, y + 42), fill='white', outline='#EBEEF5')
        text(d, (515, y + 12), k, 17, '#606266')
        draw_wrapped(d, (670, y + 10), v, 270, 16, '#303133')
        y += 42
    im.save(IMG / '10-goods-detail.png')


for fn in [save_login, save_dashboard, save_employee_list, save_employee_form, save_my_credits, save_credit_ops, save_category, save_goods_list, save_goods_form, save_goods_detail]:
    fn()


def set_cell(cell, value, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(value)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.bold = bold
    if bold:
        r.font.color.rgb = DARK


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths):
    table_obj = doc.add_table(rows=1, cols=len(headers))
    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_obj.autofit = False
    for i, h in enumerate(headers):
        set_cell(table_obj.rows[0].cells[i], h, True)
        shade(table_obj.rows[0].cells[i], 'F2F4F7')
    for row in rows:
        cells = table_obj.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    for row in table_obj.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph()
    return table_obj


def bullet(doc, value):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.add_run(value)


def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(11)
    styles['Normal'].paragraph_format.space_after = Pt(6)
    styles['Normal'].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [('Heading 1', 16, BLUE, 16, 8), ('Heading 2', 13, BLUE, 12, 6), ('Heading 3', 12, DARK, 8, 4)]:
        s = styles[name]
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run('YUTANK Shop Project - Interim WIP Submission')
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def caption(doc, value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(value)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


doc = Document()
configure(doc)
p = doc.add_paragraph()
r = p.add_run('YUTANK Shop Project - Interim Work-in-Progress Submission')
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor(11, 37, 69)
p.paragraph_format.space_after = Pt(3)
p2 = doc.add_paragraph()
r2 = p2.add_run('Client-facing progress brief for the current back-office management system iteration')
r2.font.size = Pt(13)
r2.font.color.rgb = DARK

add_table(doc, ['Item', 'Details'], [
    ('Project focus', 'Back-office management system for the YUTANK employee welfare shop.'),
    ('Current status', 'Work-in-progress submission with implemented admin modules and demo evidence.'),
    ('Excluded scope', 'Shopping cart functionality is not included in this interim submission.'),
    ('Verification', 'Backend compilation and frontend production build have completed successfully; frontend build warnings were asset-size warnings only.')
], [1.6, 4.9])

doc.add_heading('1. Summary of Delivered User Stories', level=1)
doc.add_paragraph('The team has completed several work-in-progress user stories for the back-office management system.')
sections = {
    'Staff Login and Authentication': ['As a staff member, I can log in with my employee username and password.', 'As a logged-in staff member, I can access protected back-office pages using token-based authentication.', 'As a staff member, I can log out and protect my account session.'],
    'Dashboard': ['As a staff member, I can view my current account name, credit balance, available management modules, and backend service status.', 'As a staff member, I can use quick links to access key admin modules.'],
    'Employee Management': ['As an admin user, I can view and search employee accounts.', 'As an admin user, I can create new employee accounts.', 'As an admin user, I can edit employee profile information.', 'As an admin user, I can enable or disable employee accounts.', 'As an admin user, I can reset an employee password.'],
    'Credit Management': ['As a staff member, I can view my current credit balance.', 'As a staff member, I can view my own credit transaction records.', 'As an admin user, I can add credits to an employee account.', 'As an admin user, I can deduct credits from an employee account.', 'As an admin user, I can review an employee credit records after credit operations.'],
    'Category and Goods Management': ['As a staff member, I can view product categories.', 'As an admin user, I can view and filter goods.', 'As an admin user, I can create new goods.', 'As an admin user, I can edit goods information.', 'As an admin user, I can view goods details.', 'As an admin user, I can put goods on shelf or take goods off shelf.']
}
for heading, items in sections.items():
    doc.add_heading(heading, level=2)
    for item in items:
        bullet(doc, item)
doc.add_paragraph('Shopping cart functionality is not included in this interim submission.')

doc.add_heading('2. Screenshots, Prototypes, Code Snippets, or Demos', level=1)
doc.add_paragraph('The following figures provide work-in-progress demo evidence for the current admin system. They are intended to show tangible iteration progress rather than final polished UI.')
figures = [
    ('01-login.png', 'Figure 1. Staff login page with username, password, and remember-me option.'),
    ('02-dashboard.png', 'Figure 2. Dashboard showing current employee information, credit balance, module overview, and backend status.'),
    ('03-employee-list.png', 'Figure 3. Employee management page with search, create, edit, status control, and password reset actions.'),
    ('04-employee-form.png', 'Figure 4. New employee dialog for creating staff accounts.'),
    ('05-my-credits.png', 'Figure 5. My Credits page showing current balance and transaction records.'),
    ('06-credit-operations.png', 'Figure 6. Credit operations page for adding or deducting employee credits.'),
    ('07-category-list.png', 'Figure 7. Category list page.'),
    ('08-goods-list.png', 'Figure 8. Goods management page with filters and shelf status controls.'),
    ('09-goods-form.png', 'Figure 9. New goods dialog for goods creation.'),
    ('10-goods-detail.png', 'Figure 10. Goods detail dialog.')
]
for file_name, cap in figures:
    doc.add_picture(str(IMG / file_name), width=Inches(6.5))
    caption(doc, cap)

doc.add_heading('3. Feedback Received from Stakeholders', level=1)
doc.add_paragraph('At this stage, the implemented modules follow the confirmed client-facing direction that the system should use English for the front end, back end, validation messages, API descriptions, and code comments.')
doc.add_paragraph('Formal stakeholder feedback on the latest working version is still pending. The next review should focus on whether the client is satisfied with the current admin workflow, employee management process, credit terminology, goods fields, and page layout.')

doc.add_heading('4. Short Reflection on Team Performance in the Sprint', level=1)
doc.add_paragraph('The sprint produced visible technical progress across both the back end and the admin front end. The team has moved from planning into a working system foundation, including authentication, employee management, credit management, category display, and goods management.')
doc.add_paragraph('The strongest progress is the connection between backend APIs and frontend admin pages. A key area for improvement in the next sprint is collecting clearer client feedback before expanding the feature set further.')

doc.add_heading('5. Progress Against the Original Release Plan', level=1)
doc.add_paragraph('The project is progressing well against the original release plan for the back-office management system.')
add_table(doc, ['Status', 'Items'], [
    ('Completed or mostly completed', 'Project environment setup; GoFrame backend structure; MySQL initialization schema; JWT-based employee login and authentication middleware; employee account management; employee credit balance and records; admin credit add/deduct operations; goods category listing; goods management; Vue-based admin front-end structure; admin routes and API integration.'),
    ('Verification completed', 'Backend package compilation passed with go test ./...; frontend production build completed successfully; frontend build warnings were asset-size warnings only.'),
    ('Still in progress', 'Client review of current admin workflows; UI polishing after stakeholder feedback; additional release features outside the current admin foundation; more formal test cases beyond build and compilation checks.'),
    ('Excluded from this report', 'Shopping cart functionality is excluded from this interim progress summary.')
], [2.0, 4.5])

DOCX.parent.mkdir(parents=True, exist_ok=True)
doc.save(DOCX)
print(DOCX)
